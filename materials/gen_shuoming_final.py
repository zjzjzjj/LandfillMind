# -*- coding: utf-8 -*-
"""
生成 LandfillMind 说明文档（嵌入新高精度架构图）
"""
import os
import copy
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

OUTPUT_DIR = r'E:\Son of  the SEA\260825提交包'
MATERIALS = r'f:\zj_F\LandfillMind\materials'
TEMPLATE_DIR = r'E:\下载\说明文档模版 (1)'

OUTPUT_DOCX = os.path.join(r'f:\zj_F\LandfillMind\materials', '1-说明文档_new.docx')
DIAGRAM_PATH = os.path.join(MATERIALS, 'landfillmind-architecture.png')

# ========== 工具函数 ==========
def set_run_font(run, name='微软雅黑', size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    """添加标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=14, bold=True, color=(0, 70, 127))
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        set_run_font(run, size=12, bold=True, color=(0, 100, 160))
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        set_run_font(run, size=11, bold=True, color=(51, 51, 51))
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
    return p

def add_para(doc, segments, indent=True):
    """
    添加段落，segments 是 list of (text, bold, color)
    """
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(4)
    for seg in segments:
        if isinstance(seg, str):
            run = p.add_run(seg)
            set_run_font(run)
        elif len(seg) == 2:
            text, bold = seg
            run = p.add_run(text)
            set_run_font(run, bold=bold)
        elif len(seg) == 3:
            text, bold, color = seg
            run = p.add_run(text)
            set_run_font(run, bold=bold, color=color)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.6)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_image_centered(doc, img_path, width_inches=6.2):
    """添加居中图片"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if os.path.exists(img_path):
        run.add_picture(img_path, width=Inches(width_inches))
    else:
        run.add_text(f'[图片缺失: {img_path}]')
    return p

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=9, color=(128, 128, 128))
    p.paragraph_format.space_after = Pt(8)
    return p

def fill_table_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, bg=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=10, bold=bold)
    if bg:
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), bg)
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'B0B0B0')
        borders.append(border)
    tblPr.append(borders)

# ========== 主文档生成 ==========
def build_document():
    doc = Document()
    
    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    # ===== 标题 =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('首届全国大学生人工智能+大赛')
    set_run_font(run, size=20, bold=True, color=(0, 70, 127))
    title.paragraph_format.space_after = Pt(4)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('作品说明文档')
    set_run_font(run, size=18, bold=True, color=(0, 100, 160))
    subtitle.paragraph_format.space_after = Pt(16)
    
    # 基本信息表
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    info = [
        ('作品名称', 'LandfillMind — 基于双引擎纠偏架构的生活垃圾填埋场智能诊断平台'),
        ('赛道类别', '人工智能+ 安全生产（智慧应急方向）'),
        ('团队名称', '山海之子（Son of the SEA）'),
        ('作品负责人', '郑杰'),
        ('作品版本', 'v4.2'),
        ('提交日期', '2026年8月'),
    ]
    for i, (k, v) in enumerate(info):
        fill_table_cell(table.cell(i, 0), k, bold=True, bg='E8F4FD')
        fill_table_cell(table.cell(i, 1), v, align=WD_ALIGN_PARAGRAPH.LEFT)
    for row in table.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(12)
    
    doc.add_paragraph()  # 空行
    
    # ===== 1. 作品概述 =====
    add_heading(doc, '1. 作品概述', level=1)
    add_para(doc, [
        'LandfillMind 是一款面向生活垃圾填埋场一线运维人员与安全监管人员的智能诊断平台。',
        '当前我国县级及以下填埋场普遍存在专业岩土工程师稀缺、规范查询困难、应急决策滞后等痛点，',
        '传统 LLM 应用在安全关键场景中存在数值幻觉风险，无法直接用于工程决策。'
    ])
    add_para(doc, [
        '本作品创新性地提出', ('"算—写—校"双引擎纠偏架构', True, (196, 89, 17)),
        '：以 TypeScript 确定性计算内核（14 项专业计算器、500+ 行核心公式）提供无幻觉的工程计算，',
        '以大语言模型（智谱 GLM-4-Flash 免费档）生成自然语言解释报告，再通过交叉复核层自动捕获并纠正 LLM 的数值幻觉。',
        '最终输出红/橙/黄/蓝四级风险评级、规范引用条款、完整计算书和一键应急决策包。'
    ])
    add_para(doc, [
        '系统已内置 CJJ 176-2012、GB 16889-2008、HJ 25.6-2019 等 199 条行业规范的 RAG 知识库，',
        '集成边坡稳定、渗滤液、填埋气、地下水 5 路智能体协同研判，支持 Three.js 3D 可视化与 OpenGeoSys 5.5 有限元数值模拟。',
        '全栈 TypeScript，零成本部署于 Render 免费层，形成"专业可信、零成本、开箱即用"的完整解决方案。'
    ])
    
    # ===== 2. 作品内容 =====
    add_heading(doc, '2. 作品内容', level=1)
    
    add_heading(doc, '2.1 需求分析', level=2)
    add_para(doc, [
        '我国生活垃圾填埋场超 1800 座，其中县级及以下中小型填埋场占比超过 60%。',
        '一线运维人员在面对边坡滑移、渗滤液导排堵塞、填埋气聚集、地下水污染等隐患时，',
        '面临三大核心痛点：'
    ])
    add_bullet(doc, '专业人才匮乏：岩土工程、环境工程专业工程师稀缺，现场人员难以独立完成规范查询与定量计算')
    add_bullet(doc, 'LLM 幻觉风险：直接使用通用大模型进行安全诊断，存在数值编造、规范错配等严重隐患')
    add_bullet(doc, '应急响应滞后：从发现异常到形成可执行决策文档往往需要数小时，错过黄金处置窗口')
    add_para(doc, [
        'LandfillMind 针对上述痛点，定位为"现场工程师的 AI 副驾驶"，',
        '以确定性计算保障数值可信，以大语言模型提供交互便利，以交叉复核确保安全底线。'
    ])
    
    add_heading(doc, '2.2 技术方案', level=2)
    add_para(doc, [
        'LandfillMind 采用全栈 TypeScript 技术栈，前端基于 React 18 + Vite 5 + Tailwind CSS + Three.js，',
        '后端基于 Node.js + Express，数据持久化采用 sql.js（SQLite WASM），',
        'LLM 接入智谱 GLM-4-Flash（默认免费通道）+ OpenRouter（可选并行通道），SSE 流式通信。'
    ])
    
    add_heading(doc, '2.2.1 整体架构：双引擎 + 纠偏剧场三层流水线', level=3)
    add_para(doc, [
        '系统核心架构如下图所示，可概括为"算—写—校"三层流水线，并辅以多智能体协同、专业工具集、一键应急输出与自进化反馈闭环。'
    ])
    
    # 插入架构图
    add_image_centered(doc, DIAGRAM_PATH, width_inches=6.2)
    add_caption(doc, '图 1  LandfillMind v4.2 系统架构图（核心创新：③ AI 纠偏剧场）')
    
    add_para(doc, [('引擎 A · 确定性计算内核（算）：', True),
        '由 server/calculate.ts（32KB，14 项计算器，超 500 行核心公式）负责。',
        '接收用户输入的场地监测数据（边坡高度、坡角、黏聚力、内摩擦角、渗滤液液位、CH₄ 浓度等），',
        '直接计算安全系数 Fs、产气速率、液位安全高度等关键指标，输出 100% 确定性结果，不存在任何幻觉风险。'
    ])
    add_para(doc, [('引擎 B · LLM 解释层（写）：', True),
        '由 server/llm.ts（11KB，双通道路由）负责。默认使用智谱 GLM-4-Flash 免费档，',
        '通过结构化 Prompt 工程（角色设定 + 规范注入 + 输出格式约束 + 禁止模糊措辞），',
        '将引擎 A 的计算结果转化为现场人员可读的自然语言诊断报告。'
    ])
    add_para(doc, [('引擎 C · 交叉复核纠偏层（校）—— AI 纠偏剧场：', True, (196, 89, 17)),
        '这是本作品最核心的创新点。由 server/diagnose.ts 的 runDiagnosis() 函数负责，',
        '对引擎 B 的 LLM 输出做正则 + 语义双重校验：(1) 提取 LLM 报告中声称的所有数值；',
        '(2) 与引擎 A 的真实计算值逐项比对；(3) 若偏差超过阈值（如 Fs 偏差 > 5%），自动触发纠偏，',
        '用内核真实值覆盖 LLM 编造值并添加修正标注。界面上以红/绿分屏形式展示对比（AI 报告原文 vs 内核真实结论），',
        '决策过程透明可追溯，真正做到"既利用 LLM 的语言能力，又不被其幻觉所误导"。'
    ])
    
    add_heading(doc, '2.2.2 Prompt 工程与知识库', level=3)
    add_para(doc, [
        '系统内置 199 条行业规范的结构化知识库（data/knowledge.db，352KB），通过 RAG 语义检索 + 同义词扩展（106 条映射表，如"垃圾汤"→"渗滤液"），',
        '在每次对话前动态注入最相关的规范条款。Prompt 设计采用"系统角色 + 领域约束 + 输出模板 + 数值锚定"四层结构，',
        '严格禁止使用"可能""大概"等模糊措辞，所有安全结论必须附带规范条款编号。'
    ])
    add_para(doc, [
        '此外实现了 A/B 测试框架：3 个 Prompt 变体并行运行，通过 UCB1 Bandit 算法根据用户👍/👎反馈动态调权，持续优化输出质量。'
    ])
    
    add_heading(doc, '2.2.3 多智能体协同编排', level=3)
    add_para(doc, [
        '对于复杂的多场耦合场景，系统启动 5 路专业 Agent 并行研判：边坡稳定 Agent（CJJ 176-2012）、渗滤液 Agent（GB 16889-2008 §5，HELP 模型）、',
        '填埋气 Agent（GB 16889-2008 §6，LandGEM 模型）、地下水 Agent（HJ 25.6-2019），以及综合研判 Agent（队长）。',
        '5 路 Agent 通过 LRU 缓存策略避免重复计算，最终由队长 Agent 携带 4 位同伴的结论进行冲突消解与综合研判，输出统一结论。'
    ])
    
    add_heading(doc, '2.2.4 专业工具集', level=3)
    add_bullet(doc, '计算中心：14 项独立计算器，包括边坡稳定 Fs、HELP 渗滤液产量、LandGEM 产气、HDPE 膜完整性、地基沉降、气体扩散距离等')
    add_bullet(doc, '3D 仿真：基于 Three.js 的填埋场实时几何模型，隐患点位按严重度自动着色（红/橙/黄/蓝），支持 60FPS 交互旋转与剖切')
    add_bullet(doc, 'OGS 数值模拟：内置 OpenGeoSys 5.5 FEM 有限元求解器，AI 对话可直接触发渗流/沉降/产气仿真，输出可视化曲线')
    add_bullet(doc, '专家问答：199 条规范 RAG 检索，支持自然语言提问，自动引用规范条款编号')
    
    add_heading(doc, '2.2.5 一键应急决策包', level=3)
    add_para(doc, [
        '当系统判定为红色或橙色风险时，可一键生成 A4 五节决策文档：诊断摘要、班前交底、风险矩阵、疏散路线、二维码扫码查看完整报告。',
        '从发现异常到输出可打印文档仅需约 30 秒，显著缩短应急响应时间。'
    ])
    
    add_heading(doc, '2.2.6 自进化闭环', level=3)
    add_para(doc, [
        '系统构建了完整的"用户反馈 → LLM 经验蒸馏 → 生成 KB 草案 → 专家人工审核 → 入库生效"自进化闭环。',
        '知识库从初始 38 条规范已增长至 199 条，并通过 A/B 测试持续优化 Prompt 效果。'
    ])
    
    add_heading(doc, '2.3 创新点', level=2)
    innovations = [
        ('双引擎纠偏架构（AI 纠偏剧场）：',
         '首次提出"确定性计算内核 + LLM 解释层 + 交叉复核"三层架构，通过正则+语义双重校验自动捕获并纠正 LLM 数值幻觉，',
         '以分屏可视化方式实现决策透明化，解决了 LLM 在安全关键场景中"不敢用"的核心问题。'),
        ('零成本全栈 TypeScript 方案：',
         '前后端统一语言，智谱 GLM-4-Flash 免费 LLM、Render 免费托管、UptimeRobot 免费保活，',
         '县级填埋场可零成本部署使用，极大降低了智慧安监技术的门槛。'),
        ('多智能体"5 路并行 + 队长研判"模式：',
         '针对填埋场"水-气-坡-地"多场耦合特点，设计 5 个专业 Agent 并行分析，LRU 缓存优化性能，',
         '队长 Agent 冲突消解输出统一结论，兼顾分析深度与结果一致性。'),
        ('规范知识库 + RAG + 同义词扩展：',
         '199 条规范结构化入库，106 条行业黑话同义词映射（如"垃圾汤"→渗滤液），让一线人员用口语也能获得专业级规范引用。'),
        ('一键应急决策包：',
         '30 秒生成 A4 五节可打印应急文档，含风险矩阵与疏散路线，打通从 AI 诊断到现场执行的最后一公里。'),
    ]
    for i, (title, *rest) in enumerate(innovations, 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.35
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f'创新点 {i}：{title}')
        set_run_font(run, bold=True, color=(196, 89, 17))
        run2 = p.add_run(''.join(rest))
        set_run_font(run2)
    
    add_heading(doc, '2.4 作品原型与代码', level=2)
    add_para(doc, [
        '作品已完成全功能开发并在线部署。核心代码量：前端约 8000 行（React/TSX），后端约 2000 行（Node.js/TypeScript），',
        '数据库 199 条规范知识，3D 模型与仿真模块约 1500 行。'
    ])
    
    # 代码/文件结构表
    add_para(doc, [('主要模块清单：', True)])
    mod_table = doc.add_table(rows=9, cols=3)
    mod_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(mod_table)
    headers = ['模块', '文件', '功能说明']
    for j, h in enumerate(headers):
        fill_table_cell(mod_table.cell(0, j), h, bold=True, bg='0D3B5C')
        for run in mod_table.cell(0, j).paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    modules = [
        ('计算内核', 'server/calculate.ts (32KB)', '14 项工程计算器，500+ 行核心公式'),
        ('LLM 路由', 'server/llm.ts (11KB)', '智谱/OpenRouter 双通道，SSE 流式'),
        ('纠偏诊断', 'server/diagnose.ts', 'AI 纠偏剧场核心逻辑，正则+语义复核'),
        ('多智能体', 'server/agents/', '5 路专业 Agent + 队长编排'),
        ('知识库', 'data/knowledge.db (352KB)', '199 条规范，RAG 语义检索'),
        ('前端主应用', 'src/App.tsx', 'React 主界面与路由'),
        ('3D 仿真', 'src/components/Landfill3D.tsx', 'Three.js 实时渲染'),
        ('OGS 模拟', 'server/ogs/', 'OpenGeoSys 5.5 FEM 求解器封装'),
    ]
    for i, (mod, fname, desc) in enumerate(modules, 1):
        fill_table_cell(mod_table.cell(i, 0), mod, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_table_cell(mod_table.cell(i, 1), fname, align=WD_ALIGN_PARAGRAPH.LEFT)
        fill_table_cell(mod_table.cell(i, 2), desc, align=WD_ALIGN_PARAGRAPH.LEFT)
    
    add_heading(doc, '2.5 应用场景与推广价值', level=2)
    add_para(doc, [
        'LandfillMind 可直接应用于全国 1800+ 座生活垃圾填埋场，尤其适合县级及以下缺乏专业工程师的中小型填埋场。',
        '系统零成本部署、开箱即用，一线人员经过 10 分钟培训即可上手使用。'
    ])
    add_para(doc, [
        '此外，本作品提出的"双引擎纠偏"架构具有极强的领域迁移性，可推广至矿山安全、化工应急、桥梁隧道监测等',
        '其他对数值准确性要求高的安全关键场景，具备显著的社会效益与产业价值。'
    ])
    
    # ===== 3. 其他说明 =====
    add_heading(doc, '3. 其他说明', level=1)
    add_para(doc, [
        '本作品为团队独立开发，所有代码与知识库内容均为团队原创。系统使用的开源组件包括 React、Three.js、OpenGeoSys 等，',
        '均遵循其对应开源许可证。智谱 GLM-4-Flash 为公开免费 API，Render 部署使用免费层服务，无商业 API 费用依赖。'
    ])
    add_para(doc, [
        '作品已在实际模拟场景中完成多轮测试，核心计算结果与商用岩土工程软件（如 GeoStudio）对标验证，',
        '安全系数 Fs 计算误差 < 2%，满足工程精度要求。'
    ])
    
    # ===== 页脚 =====
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('— 全文完 —')
    set_run_font(run, size=10, color=(128, 128, 128))
    
    # 保存
    doc.save(OUTPUT_DOCX)
    print(f'说明文档已保存: {OUTPUT_DOCX}')
    return OUTPUT_DOCX

if __name__ == '__main__':
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    build_document()
